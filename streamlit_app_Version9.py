import streamlit as st
from docx import Document
import re
import os
import html

LAWS_DIR = "laws"

# دالة لتوحيد النص العربي للبحث المرن
def normalize_arabic_text(text):
    text = re.sub(r'(.)\1{2,}', r'\1', text)  # إزالة المد
    text = re.sub(r'[\u064B-\u0652]', '', text)  # إزالة التشكيل
    text = re.sub('[إأآا]', 'ا', text)
    text = re.sub('[ىي]', 'ي', text)
    text = re.sub('[ة]', 'ه', text)
    text = re.sub('ؤ', 'و', text)
    text = re.sub('ئ', 'ي', text)
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub('\s+', ' ', text)
    return text.strip()

# تمييز الكلمات في النتائج
def highlight_keywords(text, keywords, normalized_keywords=None, exact_match=False):
    for kw in keywords:
        if not kw: continue
        pattern = re.compile(r"(?<!\w)"+re.escape(kw)+r"(?!\w)", re.IGNORECASE)
        text = pattern.sub(r'<mark>\g<0></mark>', text)
    if exact_match and normalized_keywords:
        normalized_text = normalize_arabic_text(text)
        for i, norm_kw in enumerate(normalized_keywords):
            if not norm_kw: continue
            original_kw = keywords[i]
            if norm_kw in normalized_arabic_text(text) and not re.search(r"(?<!\w)"+re.escape(original_kw)+r"(?!\w)", text):
                pattern = re.compile(norm_kw, re.IGNORECASE)
                def replacer(m):
                    return f'<mark class="mark-soft">{m.group(0)}</mark>'
                text = pattern.sub(replacer, text)
    return text

# حجم خط النتائج: افتراضي 18، أقصى 20، أدنى 18
if "results_font_size" not in st.session_state:
    st.session_state.results_font_size = 18

# ------------------ واجهة التبويبات ------------------
tabs = st.tabs(["🔍 البحث في القوانين", "📄 عرض القانون الكامل"])

# ------------------ تبويب البحث ------------------
with tabs[0]:
    st.markdown("""
    <style>
    mark {
        background: #ffd600 !important;
        color: #000 !important;
    }
    mark.mark-soft {
        background: #ffe082 !important;
        color: #000 !important;
    }
    .result-box-custom {
        background-color: #f1f8e9;
        color: #232323;
        padding: 20px;
        margin-bottom: 10px;
        width: 100%;
        max-width: 100%;
        border-radius: 10px;
        border: 1px solid #c5e1a5;
        direction: rtl;
        text-align: right;
        font-size: %dpx;
        transition: font-size 0.2s;
    }
    </style>
    """ % st.session_state.results_font_size, unsafe_allow_html=True)

    # أزرار تكبير وتصغير حجم خط النتائج (فقط على النتائج)
    st.write("### حجم خط مواد النتائج")
    col1, col2, col3 = st.columns([1,2,1])
    with col1:
        if st.button("➖", key="decrease_font") and st.session_state.results_font_size > 18:
            st.session_state.results_font_size -= 1
    with col2:
        st.markdown(f"<div style='text-align:center;font-size:18px;'>{st.session_state.results_font_size}px</div>", unsafe_allow_html=True)
    with col3:
        if st.button("➕", key="increase_font") and st.session_state.results_font_size < 20:
            st.session_state.results_font_size += 1

    # البحث في القوانين
    if not os.path.exists(LAWS_DIR):
        st.error(f"⚠️ مجلد '{LAWS_DIR}/' غير موجود. يرجى التأكد من وجود ملفات القوانين.")
    else:
        files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]
        if not files:
            st.warning(f"📂 لا توجد ملفات قوانين في مجلد '{LAWS_DIR}/'.")
        else:
            with st.form("main_search_form"):
                selected_file_form = st.selectbox("اختر قانونًا للبحث:", ["الكل"] + files)
                keywords_form = st.text_area("📌 اكتب كلمة أو جملة للبحث عنها:")
                article_number_input = st.text_input("أو أبحث برقم المادة:")
                exact_match = st.checkbox("تطابق تام للكلمة (لا تظهر مشتقاتها مثل تظلم عند البحث عن ظلم)")
                submitted = st.form_submit_button("🔍 بدء البحث")
            if submitted:
                results = []
                search_files = files if selected_file_form == "الكل" else [selected_file_form]
                kw_list = [k.strip() for k in keywords_form.split(",") if k.strip()] if keywords_form else []
                search_by_article = bool(article_number_input.strip())
                normalized_kw_list = [normalize_arabic_text(kw) for kw in kw_list] if kw_list else []
                norm_article = article_number_input.strip()
                for file in search_files:
                    doc = Document(os.path.join(LAWS_DIR, file))
                    law_name = file.replace(".docx", "")
                    last_article = "غير معروفة"
                    current_article_paragraphs = []
                    for para in doc.paragraphs:
                        txt = para.text.strip()
                        if not txt:
                            continue
                        match = re.match(r"مادة\s*[\(]?\s*(\d+)[\)]?", txt)
                        if match:
                            if current_article_paragraphs:
                                full_text = "\n".join(current_article_paragraphs)
                                add_result = False
                                simple_full_text = normalize_arabic_text(full_text)
                                if search_by_article and last_article == norm_article:
                                    add_result = True
                                elif normalized_kw_list:
                                    for idx, kw in enumerate(normalized_kw_list):
                                        if not kw:
                                            continue
                                        if exact_match:
                                            pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                            if re.search(pattern, simple_full_text):
                                                add_result = True
                                                break
                                        else:
                                            if kw in simple_full_text:
                                                add_result = True
                                                break
                                if add_result:
                                    highlighted = highlight_keywords(full_text, kw_list, normalized_keywords=normalized_kw_list, exact_match=exact_match) if kw_list else full_text
                                    results.append({
                                        "law": law_name,
                                        "num": last_article,
                                        "text": highlighted,
                                        "plain": full_text
                                    })
                                current_article_paragraphs = []
                            last_article = match.group(1)
                        current_article_paragraphs.append(txt)
                    if current_article_paragraphs:
                        full_text = "\n".join(current_article_paragraphs)
                        add_result = False
                        simple_full_text = normalize_arabic_text(full_text)
                        if search_by_article and last_article == norm_article:
                            add_result = True
                        elif normalized_kw_list:
                            for idx, kw in enumerate(normalized_kw_list):
                                if not kw:
                                    continue
                                if exact_match:
                                    pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                    if re.search(pattern, simple_full_text):
                                        add_result = True
                                        break
                                else:
                                    if kw in simple_full_text:
                                        add_result = True
                                        break
                        if add_result:
                            highlighted = highlight_keywords(full_text, kw_list, normalized_keywords=normalized_kw_list, exact_match=exact_match) if kw_list else full_text
                            results.append({
                                "law": law_name,
                                "num": last_article,
                                "text": highlighted,
                                "plain": full_text
                            })
                if results:
                    st.success(f"تم العثور على {len(results)} نتيجة.")
                    for i, r in enumerate(results):
                        with st.expander(f"📚 المادة ({r['num']}) من قانون {r['law']}", expanded=True):
                            st.markdown(f'''
                            <div class="result-box-custom">
                                <p style="line-height:1.8;margin-top:0px;">
                                    {r["text"]}
                                </p>
                            </div>
                            ''', unsafe_allow_html=True)
                            st.code(r["plain"], language="text")
                else:
                    st.info("لم يتم العثور على نتائج مطابقة للبحث.")

# ------------------ تبويب عرض القانون الكامل ------------------
with tabs[1]:
    st.markdown("<h3 style='text-align:center;'>عرض القانون الكامل</h3>", unsafe_allow_html=True)
    if not os.path.exists(LAWS_DIR):
        st.error(f"⚠️ مجلد '{LAWS_DIR}/' غير موجود. يرجى التأكد من وجود ملفات القوانين.")
    else:
        files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]
        if not files:
            st.warning(f"📂 لا توجد ملفات قوانين في مجلد '{LAWS_DIR}/'.")
        else:
            law_sel = st.selectbox("اختر القانون:", files, key="law_select_for_view")
            if law_sel:
                doc = Document(os.path.join(LAWS_DIR, law_sel))
                st.markdown(f"<h5 style='text-align:center;color:#1976d2'>{law_sel.replace('.docx','')}</h5>", unsafe_allow_html=True)
                law_text = ""
                for para in doc.paragraphs:
                    txt = para.text.strip()
                    if txt:
                        law_text += txt + "\n\n"
                st.text_area("القانون كامل:", law_text, height=700, key="full_law_view_text", disabled=True)